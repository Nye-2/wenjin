#!/usr/bin/env python3
# -*- coding: utf-8 -*-

from __future__ import annotations

from dataclasses import dataclass
from pathlib import Path

from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml.ns import qn
from docx.shared import Cm, Pt


ROOT = Path("/home/cjz")
TEMPLATE = ROOT / "cjz/2026届毕业设计（论文）撰写格式（计算机学院模板）/1-毕业设计论文封面（无校外导师）.docx"
OUT = ROOT / "cjz/2026届毕业设计（论文）撰写格式（计算机学院模板）/问津_毕业论文初稿_super_agent_harness_架构优化版.docx"


@dataclass(frozen=True)
class CoverMeta:
    title: str
    college: str
    major: str
    class_name: str
    student_id: str
    student_name: str
    advisor: str
    advisor_title: str
    completion_date: str


COVER = CoverMeta(
    title="基于Super Agent Harness的问津学术写作工作台设计与实现",
    college="计算机科学与技术学院",
    major="计算机科学与技术",
    class_name="待填写",
    student_id="待填写",
    student_name="待填写",
    advisor="待填写",
    advisor_title="待填写",
    completion_date="2026年4月8日",
)


KEYWORDS_ZH = "Super Agent Harness；学术写作工作台；长期记忆系统；多智能体协同；系统架构设计"
KEYWORDS_EN = "Super Agent Harness; academic writing workbench; long-term memory system; multi-agent collaboration; architecture design"


ABSTRACT_ZH = [
    "随着生成式人工智能从通用问答逐步走向复杂知识工作，面向毕业论文、SCI 论文与研究申报场景的智能写作平台不再只是一个“文本生成器”，而需要具备工作区组织、上下文维持、任务调度、过程追踪与最终交付等完整能力。现有多数系统虽然在选题、摘要、润色或文献问答等局部任务中已经能够提供帮助，但其底层往往缺少统一编排框架，导致各模块之间上下文割裂、状态不可恢复、结果难以复用，难以支撑论文写作的长链路协作。针对这一问题，本文围绕问津学术写作工作台，提出并实现了一套面向学术写作场景的 Super Agent Harness。",
    "本文将 Super Agent Harness 定义为：以 Lead Agent 为统一入口，以中间件链为约束与注入手段，以 Subagent Manager 和 ParallelExecutor 为任务拆解与并行执行核心，以 Task Runtime 为状态承载机制，以 Artifact 和 Activity 为结果沉淀接口，以 Memory System 为跨轮次个性化上下文来源的统一代理底座。围绕这一底座，系统进一步建立了 Router、Application Handler、Task Handler、Feature Registry 和 Contracts 的清晰平台边界，并通过 canonical route、canonical payload 与 canonical result contract 保证多类型工作区共用一套执行主链路。",
    "在关键机制上，本文重点完成了三类设计。第一，设计了基于 ThreadState 的代理状态模型和中间件链模型，将 workspace_context、literature_context、knowledge_context、memory_context 等要素显式挂载到统一状态中。第二，设计并实现了长期记忆系统，包括会话后异步知识抽取、按 workspace 与相似度排序的记忆注入、基于 token 预算的上下文压缩，以及超阈值后的 memory compaction，从而支持长期写作任务的连续性。第三，设计了子代理生命周期管理与活动投影机制，使 feature task、chat thread、subagent task 与 artifact creation 能够汇聚为统一时间线，提高系统的可观测性和可解释性。",
    "问津当前支持 thesis、sci、proposal、software_copyright 和 patent 五类工作区、二十项 feature，并在毕业论文工作区中提供深度调研、文献管理、开题调研、论文写作、图表生成与编译导出六大模块。系统前端采用 Next.js，后端采用 FastAPI，结合 PostgreSQL、Redis、LangGraph 与 Docker 化 LaTeX 编译链构成完整闭环。本文实际运行 `uv run pytest tests/services/test_workspace_activity_service.py -q` 与 `uv run pytest tests/task/test_thesis_writing_end_to_end.py -q` 两组测试，结果分别为 7 项全部通过和 3 项全部通过。综合来看，本文实现的问津平台证明了以 Super Agent Harness 为中心、以 Memory System 为长期上下文支撑的学术写作架构具有较强的工程可行性，也为后续扩展更细粒度的评审、引用治理、成本控制与跨模型编排奠定了基础。",
]


ABSTRACT_EN = [
    "As generative AI evolves from generic dialogue into complex knowledge work, academic writing platforms for thesis, journal paper, and proposal workflows can no longer be treated as simple text generators. They must support workspace organization, long-context continuity, task scheduling, process observability, and final delivery. Many existing tools already assist with isolated tasks such as topic brainstorming, abstract generation, polishing, or literature Q&A, but they usually lack a unified orchestration substrate. As a result, context becomes fragmented across modules, state is difficult to recover, and intermediate outputs are hard to reuse in long writing workflows. To address this problem, this thesis designs and implements Wenjin, an academic writing workbench centered on a Super Agent Harness.",
    "In this thesis, the Super Agent Harness is defined as a unified orchestration substrate that combines a lead agent as the single entry point, middleware chains for contextual injection and runtime constraints, a subagent manager and parallel executor for task decomposition, a task runtime for state persistence, artifact and activity projections for result tracking, and a memory system for long-term personalized context. On top of this substrate, the platform establishes explicit layer boundaries among routers, application handlers, task handlers, feature registry, and contracts, while enforcing canonical routes, canonical payloads, and canonical result contracts.",
    "The thesis emphasizes three architectural mechanisms. First, it introduces a typed thread state and middleware-driven context model that explicitly carries workspace, literature, knowledge, and memory contexts. Second, it implements a long-term memory system with asynchronous conversation capture, knowledge extraction, similarity-based memory ranking, token-budgeted prompt injection, and memory compaction. Third, it builds durable subagent lifecycle management and unified workspace activity projection so that feature tasks, chat threads, subagent tasks, and artifacts can be observed through a single timeline. Built with Next.js, FastAPI, PostgreSQL, Redis, LangGraph, and Docker-based LaTeX compilation, Wenjin currently supports five workspace categories and twenty features. Local automated verification shows that the workspace activity service tests passed 7 out of 7 cases and the thesis writing end-to-end tests passed 3 out of 3 cases. These results demonstrate that a Super Agent Harness architecture with a dedicated memory system is feasible and effective for academic writing platforms.",
]


TOC_ITEMS = [
    "摘    要",
    "ABSTRACT",
    "第1章 绪论",
    "第2章 相关技术与理论基础",
    "第3章 系统需求与总体架构设计",
    "第4章 Super Agent Harness与Memory系统设计",
    "第5章 核心功能模块与界面实现",
    "第6章 系统测试与分析",
    "第7章 总结与展望",
    "参考文献",
    "致谢",
]


BODY = [
    {
        "title": "第1章 绪论",
        "sections": [
            {
                "title": "1.1 研究背景与问题提出",
                "paragraphs": [
                    "在本科毕业论文写作过程中，学生往往需要在项目代码、文献资料、导师意见、章节草稿、图表资源和最终版式之间频繁切换。过去，这些环节通常由多个独立工具分别承担：文献检索依赖搜索平台，正文撰写依赖文档编辑器，排版与导出依赖 LaTeX 工程，沟通记录和写作中间产物则散落在聊天软件或个人笔记中。工具链虽然丰富，但流程之间并不连贯，用户很难让“上一轮的研究结论”稳定成为“下一轮的写作上下文”。",
                    "生成式人工智能为上述问题提供了新的可能。大模型已经具备结构化表达、检索增强、工具调用和多阶段推理能力，因此学术写作平台有机会从单纯的“辅助生成器”升级为“研究工作台”。然而，如果平台只是在传统 Web 应用上简单叠加若干 AI 按钮，那么最终得到的仍然只是离散能力集合，而非真正的写作系统。尤其在毕业论文场景中，用户更需要的是跨阶段连续性、任务可追踪性、结果可回放性以及最终交付闭环。",
                    "因此，本文将研究问题聚焦为：如何构建一套能够长期支撑学术写作流程的统一代理底座，使大模型能力不再以零散功能存在，而是以稳定的软件系统方式组织起来。围绕这一问题，本文提出以 Super Agent Harness 为核心的系统架构，并在问津平台中完成工程实现。",
                ],
            },
            {
                "title": "1.2 国内外研究与实践现状",
                "paragraphs": [
                    "从方法层面看，近年来与智能体相关的研究主要集中在推理增强、检索增强和工具增强三个方向。ReAct 通过交替执行 reasoning 与 acting，使模型在复杂任务中可以边思考边调用外部工具[1]；RAG 通过把外部检索结果引入生成过程，降低模型对参数知识的单点依赖[2]；Chain-of-Thought 提示则证明了显式中间推理过程能够增强复杂问题求解能力[3]。这些方法为学术写作系统的代理化演进提供了理论基础。",
                    "从产品层面看，国内外已经出现了大量与论文辅助写作相关的系统，例如文献问答平台、论文润色工具、参考文献管理工具和整稿写作助手等。这些系统大多能够在某一特定环节提高效率，但经常存在两个共性问题。第一，系统只聚焦某个单功能点，无法自然承接前后步骤。第二，系统对状态管理、任务生命周期和中间产物缺少统一治理，一旦用户需要跨轮次延续上下文，系统就会退化成新的空白对话。",
                    "围绕复杂代理编排的工程实践也在持续演进。LangGraph 等框架开始强调显式状态流、阶段节点和执行图，这为任务拆解、异常恢复和过程观测提供了较好的基础[9]。但框架本身并不直接等于产品架构。对于学术写作平台而言，还需要进一步解决工作区边界、权限绑定、长期记忆、任务投影、交付物管理等问题。本文的意义就在于将这些问题统一纳入同一套平台架构之中。",
                ],
            },
            {
                "title": "1.3 研究内容、目标与论文结构",
                "paragraphs": [
                    "本文的研究目标不是简单实现一个“可以生成论文”的原型系统，而是设计一套能够支持长链路学术写作的工程化架构。具体而言，本文围绕问津项目完成了四项核心工作：一是分析学术写作平台的业务需求与非功能需求，明确为何需要统一编排底座；二是提出并实现以 Super Agent Harness 为中心的平台架构；三是围绕 Memory System、Subagent Orchestration、Task Runtime、Artifact/Activity Projection 等关键机制进行设计；四是结合论文写作、深度调研和 LaTeX 主稿等功能模块验证架构的实际价值。",
                    "全文结构安排如下：第一章介绍研究背景、现状和论文整体思路；第二章介绍智能体、记忆系统、状态图编排和异步任务等相关技术；第三章从需求分析出发给出系统总体架构设计；第四章聚焦 Super Agent Harness 与 Memory System 的实现机制；第五章说明核心功能模块与前端工作台界面；第六章给出测试结果与架构分析；第七章进行总结与展望。",
                ],
            },
        ],
    },
    {
        "title": "第2章 相关技术与理论基础",
        "sections": [
            {
                "title": "2.1 大模型智能体与Harness化思路",
                "paragraphs": [
                    "传统的 AI 应用通常以“输入一段文本，返回一段文本”为基本交互模式，而复杂学术任务并不能被压缩到如此简单的 I/O 结构。论文写作涉及问题理解、材料调用、上下文延续、任务拆分、结果校验和长期演进等多个步骤，因此系统核心不应只是一个模型接口，而应是一套能稳定组织代理行为的 Harness。Harness 的本质，是把模型、工具、状态、限制条件和产物接口收束到同一运行时之中。",
                    "在本文语境下，Super Agent Harness 不是多智能体数量上的“多”，而是控制面上的“统一”。它强调单入口、单状态面、单契约和单投影面。用户可以从一个工作区对话入口出发，但系统内部可以调用中间件、feature graph、subagent、runtime block 与 artifact 存储。这样的设计既保留了复杂能力，又避免把平台做成难以维护的脚本堆栈。",
                ],
            },
            {
                "title": "2.2 ReAct、RAG与状态图编排",
                "paragraphs": [
                    "ReAct 机制为复杂代理任务提供了“思考-行动-再思考”的组织方式，特别适合学术写作场景中先分析问题、再决定是否检索文献或调用工具的需求[1]。RAG 机制则通过引入外部知识源，为模型生成提供事实支撑，使系统能够把工作区内文献、知识面板和项目上下文真正纳入推理过程[2]。",
                    "与单轮对话不同，学术写作中的很多任务天然具有阶段结构，例如深度调研需要先发现文献，再识别空白，再综合结论；论文写作需要先生成大纲，再写章节，再改写和复核。状态图编排的优势在于，它可以显式表达节点之间的依赖关系、并发关系和结果汇总逻辑。问津基于 LangGraph 的思想进一步构建了 workspace feature graph，使 feature 执行从请求入口到结果落库形成统一链路。",
                ],
            },
            {
                "title": "2.3 长期记忆与上下文连续性",
                "paragraphs": [
                    "在长周期学术任务中，上下文连续性是决定系统可用性的关键因素。仅靠对话窗口中的最近若干轮消息，无法稳定承载用户偏好、研究背景、当前阶段、导师要求和长期目标等信息。因此，学术写作平台需要在短期线程状态之外，引入长期记忆系统。长期记忆并不意味着把所有历史都塞进提示词，而是从历史交互中提炼对未来协作有价值的信息，并在合适时机注入当前推理过程。",
                    "问津的记忆设计正体现了这一思想。系统并不直接复用全量聊天记录，而是将历史内容抽取为 preference、knowledge、context、behavior 和 goal 五类知识条目，并结合置信度、workspace 匹配关系和当前上下文相似度进行排序。随后，系统再在 max_injection_tokens 预算内格式化为 memory_context 注入提示词。这样的策略在保证连续性的同时，避免了无边界的上下文膨胀。",
                ],
            },
            {
                "title": "2.4 异步任务、事件流与活动投影",
                "paragraphs": [
                    "深度调研、论文整稿写作和 LaTeX 编译等功能都不是适合同步返回的轻量操作。对于这类任务，平台必须通过任务运行时管理生命周期，并持续向前端反馈状态。问津为此设计了 TaskRecord 与 runtime_state 结构，并通过任务状态接口和事件流把进度、阶段和结果暴露给前端。这样，系统中的“AI 正在做什么”不再是不可见的黑箱，而成为可显示、可刷新、可中断的对象。",
                    "更进一步，问津并没有让任务状态只存在于底层存储中，而是通过 WorkspaceActivityService 把 task、chat、artifact 和 subagent 统一映射为工作区活动流。这使系统具备更好的过程解释能力，也为论文写作的版本回顾和历史追踪提供了基础。",
                ],
            },
            {
                "title": "2.5 前后端分离与可交付主稿链路",
                "paragraphs": [
                    "一个完整的学术写作平台不仅要“会生成内容”，还要“能交付结果”。问津采用 Next.js 构建前端工作台，采用 FastAPI 构建网关与业务服务[4][5]，以 PostgreSQL 存储结构化实体，以 Redis 提供运行时辅助能力[6][7][8]。在最终交付层面，系统进一步引入 LaTeXProjectService 和 CompileService，使论文正文、项目文件和 PDF 导出形成闭环。这一点保证了平台不仅服务于对话生成，也服务于毕业设计提交这一最终目标。",
                ],
            },
        ],
    },
    {
        "title": "第3章 系统需求与总体架构设计",
        "sections": [
            {
                "title": "3.1 业务需求与非功能需求分析",
                "paragraphs": [
                    "从业务目标看，问津面向的不是单一论文类型，而是 thesis、sci、proposal、software_copyright 与 patent 五类工作区。因此，平台必须支持多类型 workspace 的统一组织和差异化 feature 扩展。对于毕业论文工作区，系统至少应支持深度调研、文献管理、开题调研、论文写作、图表生成和编译导出六类能力，并允许这些能力在同一工作区中按阶段串联。",
                    "从非功能目标看，平台必须满足可扩展、可追踪、可恢复和可测试四项要求。可扩展意味着新增 feature 时应先在 registry 中声明，再接入 graph、service 和前端，而非复制一整套流程。可追踪意味着任务、子代理与产物都应有统一投影。可恢复意味着线程状态和长期记忆不能完全依赖单次上下文。可测试意味着关键执行结果必须有稳定数据契约作为断言基础。正是这些非功能需求，推动平台采用 Super Agent Harness 而不是简单模块拼装。",
                ],
            },
            {
                "title": "3.2 平台分层边界设计",
                "paragraphs": [
                    "问津在架构设计中明确区分 Router、Application Handler、Task Handler、Feature Registry 和 Contracts 五类平台边界。Router 只负责 HTTP 协议适配、鉴权依赖注入和请求校验；Application Handler 负责权限、积分、幂等和任务提交等跨服务编排；Task Handler 负责异步执行、进度上报和 LangGraph 调用；Feature Registry 负责 workspace feature 的单一事实源；Contracts 负责跨层共享的数据结构。这一边界设计避免了路由层业务回流和任务层协议污染，使系统更适合持续演进。",
                    "同时，平台将 `/api/workspaces/{workspace_id}/features/{feature_id}/execute` 作为 feature 的 canonical 执行入口，将 `/api/tasks/*` 作为统一任务状态读取面，将 `/api/workspaces/{workspace_id}/artifacts*` 作为统一成果读写入口。统一路由面的意义在于，前端不需要为不同能力记忆不同入口，后端也能围绕单一主链路做权限控制、测试和可观测性设计。",
                ],
                "table": {
                    "title": "表3-1 平台分层职责划分",
                    "headers": ["层级", "主要职责", "禁止事项"],
                    "rows": [
                        ["Router", "HTTP 协议适配、鉴权注入、参数校验", "直接处理积分策略、直接编排复杂流程"],
                        ["Application Handler", "跨服务编排、权限校验、任务提交", "处理 HTTP 响应细节"],
                        ["Task Handler", "异步执行、进度更新、调用 LangGraph", "处理路由层协议与前端格式"],
                        ["Feature Registry", "workspace feature 元信息单一事实源", "编写业务流程分支逻辑"],
                        ["Contracts", "跨层数据结构与序列化约束", "直接执行业务逻辑"],
                    ],
                },
            },
            {
                "title": "3.3 总体架构设计",
                "paragraphs": [
                    "系统总体上可以抽象为五层。第一层是前端工作台层，负责提供首页、工作区主页、聊天界面、活动时间线、知识面板和 LaTeX 主稿台。第二层是网关与应用层，负责路由管理、请求校验、工作区访问控制和任务提交。第三层是本文的核心，即 Super Agent Harness 层，包括 Lead Agent、中间件链、Subagent Manager、ParallelExecutor、feature graph 与 Memory System。第四层是业务服务层，包括深度调研、论文写作、Dashboard、WorkspaceActivity、LaTeX 项目与编译等服务。第五层是数据与运行时层，主要由 PostgreSQL、Redis、文件系统和 Docker 编译环境组成。",
                    "与普通前后端分离项目相比，问津最关键的差异并不在于使用了多少技术组件，而在于把“代理编排层”作为独立的一层对待。该层既不是简单的 SDK 封装，也不是某个 feature 的局部逻辑，而是整个系统的控制中心。其输入来自工作区、线程、模型配置和用户上下文，其输出则被统一归约为 task、artifact、activity 和 workspace refresh 事件。这使系统在面对多模块扩展时仍能保持一致的控制逻辑。",
                ],
                "figures": [
                    {
                        "caption": "图3-1 系统总体架构图占位",
                        "hint": "建议插入一张分层架构图：前端工作台 -> 网关/应用层 -> Super Agent Harness -> 业务服务层 -> 数据/运行时层。",
                    }
                ],
            },
            {
                "title": "3.4 数据模型与状态模型设计",
                "paragraphs": [
                    "问津通过结构化数据模型支撑长链路学术任务。Workspace 用于组织项目边界；ChatThread 用于保存会话级状态与消息；TaskRecord 用于保存异步任务生命周期和 runtime_state；Artifact 用于沉淀研究报告、章节草稿和编译结果，并通过 parent_artifact_id 形成 lineage；SubagentTaskRecord 用于持久化子代理生命周期；UserKnowledge 用于存储长期记忆条目；LatexProject 用于描述 LaTeX 主稿工程。上述结构共同构成了平台的“持久化背板”。",
                    "除数据库实体之外，ThreadState 也是一个关键设计。系统并没有把代理上下文塞进零散的局部变量，而是把 messages、workspace_id、workspace_type、literature_context、knowledge_context、memory_context、cited_papers 和 subagent_tasks 等字段显式挂载在统一状态对象上。这种做法提高了代理运行的透明度，也使不同中间件和 graph 节点之间可以围绕同一状态接口协作。",
                ],
                "table": {
                    "title": "表3-2 核心数据与状态对象",
                    "headers": ["对象", "关键字段", "作用"],
                    "rows": [
                        ["Workspace", "name、type、discipline、config", "定义工作区边界与学术任务类型"],
                        ["ChatThread", "workspace_id、model、skill、messages", "保存工作区主线程会话"],
                        ["TaskRecord", "task_type、status、payload、result、runtime_state", "承载异步任务生命周期"],
                        ["Artifact", "type、content、parent_artifact_id、version", "沉淀与追踪中间产物"],
                        ["SubagentTaskRecord", "thread_id、subagent_type、status、output_preview", "保存子代理生命周期"],
                        ["UserKnowledge", "category、content、confidence、workspace_context", "存储长期记忆条目"],
                        ["ThreadState", "memory_context、cited_papers、subagent_tasks", "统一代理运行状态接口"],
                    ],
                },
            },
            {
                "title": "3.5 统一执行链路设计",
                "paragraphs": [
                    "问津将 workspace feature 的执行统一收敛到一条标准链路中：前端从 canonical route 发起 feature execute 请求，Router 接收请求并注入当前用户，Application Handler 完成 workspace ownership、feature lookup、credit/policy check 和 task submission，TaskService 创建任务后交由 Task Handler 异步执行，Task Handler 再调用 `workspace_lead_agent.execute_feature_graph(...)` 完成 graph 运行与结果封装，最终将 artifact、runtime state 和 workspace refresh 事件返回前端。",
                    "这种执行链路的意义在于，它把“复杂学术能力”转化为“统一工程接口”。无论是深度调研、论文写作还是图表生成，系统最终都围绕相同的 payload 规范和 result contract 运行，从而减少 feature 之间的实现分叉。这也是后续 Super Agent Harness 能够成为平台底座的前提。",
                ],
                "figures": [
                    {
                        "caption": "图3-2 workspace feature 执行时序图占位",
                        "hint": "建议插入一张时序图：Frontend -> Router -> FeatureExecutionHandler -> TaskService -> TaskHandler -> workspace_lead_agent -> artifacts/runtime/activity。",
                    }
                ],
            },
        ],
    },
    {
        "title": "第4章 Super Agent Harness与Memory系统设计",
        "sections": [
            {
                "title": "4.1 Harness设计目标与核心抽象",
                "paragraphs": [
                    "问津中的 Super Agent Harness 面向的不是单轮生成，而是面向工作区级别的长期学术协作。因此，其设计目标主要有四项：第一，保持统一代理入口，使所有复杂能力都可以从工作区对话或 feature 调用进入；第二，显式管理上下文，使文献、知识、记忆和工作区状态能够稳定注入代理推理；第三，显式管理执行生命周期，使任务、子代理、运行时块和最终结果都可被追踪；第四，显式管理沉淀物，使研究报告、章节草稿和编译结果都能被复用。",
                    "围绕这些目标，Harness 形成了五个核心抽象：Lead Agent、Middleware Chain、Subagent Runtime、Task Runtime、Artifact/Activity Projection。Lead Agent 负责统一决策入口；Middleware Chain 负责把环境约束与上下文注入显式化；Subagent Runtime 负责任务拆分与并发治理；Task Runtime 负责 feature 执行过程中的状态块与阶段进度；Artifact/Activity Projection 则负责把结果沉淀为可见对象。本文认为，这五类抽象共同构成了平台的系统核心。",
                ],
            },
            {
                "title": "4.2 Lead Agent与中间件链设计",
                "paragraphs": [
                    "Lead Agent 工厂的职责并不是简单创建模型实例，而是对运行时配置进行规范化，并按 workspace 类型、模型能力和当前上下文装配合适工具与中间件。系统会在 configurable 中补齐 `model_name`、`supports_vision` 与 `subagent_enabled` 等字段，再结合工作区类型生成相应系统提示，从而把“模型能不能做这件事”和“当前工作区为什么要做这件事”统一起来。",
                    "更关键的是中间件链。问津已接入 WorkspaceContext、LiteratureContext、KnowledgeContext、CitationContext、Memory、Clarification、Sandbox、Summarization、Uploads、ViewImage、TodoList 与 SubagentLimit 等中间件。这些中间件并不是装饰性的前后置钩子，而是实际参与上下文构建与行为约束的运行时组件。例如，WorkspaceContext 负责注入项目语境，LiteratureContext 负责注入工作区文献，MemoryMiddleware 负责长期记忆注入与会话后捕获，SandboxMiddleware 负责执行环境隔离，SubagentLimitMiddleware 负责约束子代理数量。",
                    "中间件链的重要性在于，它把原本可能散落在 prompt 中的隐式规则提升为可维护的软件边界。这种设计使代理不再完全依赖大段自然语言提示词保持稳定，而是依赖结构化中间件共同塑形。对于学术写作这类高上下文、强流程、需要谨慎约束的场景而言，这种做法比“单提示词大一统”更稳妥。",
                ],
                "table": {
                    "title": "表4-1 Harness核心中间件职责",
                    "headers": ["中间件", "主要作用", "对应价值"],
                    "rows": [
                        ["WorkspaceContext", "注入工作区名称、类型、学科与配置", "让代理理解当前任务边界"],
                        ["LiteratureContext", "注入工作区文献上下文", "提升回答的事实性与相关性"],
                        ["KnowledgeContext", "注入知识面板与产物上下文", "增强跨模块复用能力"],
                        ["MemoryMiddleware", "长期记忆注入与会话后捕获", "维持跨轮次连续性"],
                        ["Clarification", "识别模糊输入并缩窄问题空间", "降低错误执行概率"],
                        ["Sandbox", "限制文件与执行环境边界", "保证系统安全性"],
                        ["SubagentLimit", "控制子代理开启数量", "避免并发失控"],
                        ["Summarization", "压缩长对话状态", "缓解上下文膨胀问题"],
                    ],
                },
            },
            {
                "title": "4.3 Memory系统设计",
                "paragraphs": [
                    "Memory System 是本文架构中的另一个重点。其作用并不是替代数据库，也不是保存所有聊天内容，而是提炼“对未来协作有价值的长期信息”。在问津中，记忆条目被组织为 preference、knowledge、context、behavior 和 goal 五类，并由 `UserKnowledge` 模型持久化存储。每条记忆都带有 confidence、source 和 workspace_context 字段，这意味着系统既能处理全局偏好，也能处理工作区特定阶段的信息。",
                    "在注入阶段，MemoryMiddleware 会在 `before_model` 钩子中读取当前用户与工作区上下文，先从 LRU 缓存中尝试命中 memory_context，若未命中，则根据最近若干轮 Human/AI 消息生成 conversation_context，再调用 `build_memory_context(...)` 加载和排序记忆。排序阶段结合 workspace 匹配、当前上下文相似度与置信度进行综合打分，之后再根据 `max_injection_tokens` 预算格式化为 `<academic_memory>` 块写回 ThreadState。这使 memory 不会无边界膨胀，而是以“最相关、最有价值”的形式参与推理。",
                    "在捕获阶段，MemoryMiddleware 会在 `after_model` 钩子中过滤出有效的 user/assistant 轮次，将其通过 `MemoryQueue` 以 30 秒 debounce 策略异步合并，再调用 `extract_and_persist_knowledge(...)` 完成知识抽取。抽取结果只有在置信度达到阈值时才会持久化，并在超出 `max_facts` 上限后触发 `_maybe_compact_memory(...)`，进一步调用 `compact_user_memory(...)` 对记忆进行合并、去重与压缩。也就是说，问津的记忆系统并非简单追加，而是具备完整的注入、捕获、筛选、压缩和再利用闭环。",
                    "从系统设计角度看，Memory System 的价值有两点。第一，它显著提升了长期写作场景中的连续性，使用户不需要在每次对话中重复解释偏好、背景与目标。第二，它把“上下文管理”从 prompt 工程问题提升为系统工程问题，使记忆成为平台一级能力，而不是附属技巧。",
                ],
                "table": {
                    "title": "表4-2 Memory系统关键机制",
                    "headers": ["环节", "实现机制", "设计目的"],
                    "rows": [
                        ["会话捕获", "filter_messages_for_memory + MemoryQueue debounce", "降低频繁写库和噪声输入"],
                        ["知识抽取", "extract_and_persist_knowledge", "从对话中沉淀长期有价值信息"],
                        ["记忆注入", "build_memory_context + format_knowledge_for_prompt", "把记忆转化为可控提示词块"],
                        ["相关性排序", "workspace 匹配 + similarity + confidence", "优先注入最有价值记忆"],
                        ["容量治理", "max_facts + compact_user_memory", "防止记忆膨胀与重复累积"],
                    ],
                },
                "figures": [
                    {
                        "caption": "图4-1 Memory系统流程图占位",
                        "hint": "建议插入流程图：消息过滤 -> debounce 队列 -> 知识抽取 -> UserKnowledge 持久化 -> build_memory_context -> memory_context 注入 ThreadState。",
                    }
                ],
            },
            {
                "title": "4.4 Subagent调度与并行执行设计",
                "paragraphs": [
                    "在复杂 feature 中，主代理不应亲自完成所有子任务，而应将局部职责委派给更专门的 subagent。问津的 GlobalSubagentManager 负责这一过程。它以 thread_id 为线程级上下文主键，维护 owner_user_id、workspace_id、任务定义、任务结果和活动任务集合，使子代理始终处在可追踪、可取消、可归属的范围内。与此同时，`build_subagent_task(...)` 与 `build_subagent_metadata(...)` 会对 max_turns、timeout、tools、model_name、workspace_id 和 user_id 等参数做统一规范化，为后续权限校验与结果持久化打下基础。",
                    "在并行层面，ParallelExecutor 使用 ExecutionPhase、PhasedPlan 和 PhaseResult 对复杂任务的阶段结构进行编码，并通过信号量与 fail-fast 策略控制并发风险。这样，系统可以把“适合并行”的部分并行化，把“必须依赖上一步”的部分顺序化。深度调研模块即采用了此设计：文献发现并行、研究空白顺序化、综合与交叉校验顺序化。该机制让子代理协同从“多开几个任务”升级为“有编排的多智能体执行”。",
                ],
            },
            {
                "title": "4.5 Task Runtime与Canonical Result Contract设计",
                "paragraphs": [
                    "Super Agent Harness 能够长期稳定运行的另一个关键，在于任务运行时与结果契约。问津把 feature 执行统一抽象为 task，并通过 TaskService 管理提交、状态、并发限制和事件广播。TaskHandler 再负责绑定 progress tracker、runtime state 和 workspace feature runtime block，将 graph 执行过程中的当前阶段、运行信息和结果产物持续回写给前端。",
                    "在结果层，系统明确要求 workspace feature 任务遵循统一结构，关键字段包括 `success`、`feature_id`、`feature_name`、`workspace_type`、`handler_key`、`message`、`data`、`artifacts` 和 `refresh_targets`。这种 canonical result contract 的意义在于，前端不必为不同 feature 分别发明新的消费方式，测试层也可以围绕统一结构做断言。对平台演进而言，这种统一契约比单个 feature 做得多漂亮更重要。",
                ],
            },
            {
                "title": "4.6 Artifact、Activity与Summary投影设计",
                "paragraphs": [
                    "问津并未把执行结果局限于最终答案文本，而是通过 Artifact、Activity 和 Summary 三个投影面把过程信息暴露给用户。Artifact 用于保存调研报告、章节草稿、图表与编译结果，并通过 lineage 支持父子追踪；Activity 用于把 feature task、chat thread、subagent task 和 artifact creation 聚合为统一工作区时间线；Summary 则通过 WorkspaceSummaryService 对当前阶段、下一步建议、风险项和近期活动进行聚合。这三个投影面共同构成了系统的人机协作界面。",
                    "尤其值得强调的是 Activity 设计。WorkspaceActivityService 会分别读取 TaskRecord、ChatThread、Artifact 和 SubagentTaskRecord，并将其转换为统一的 canonical activity item。这样，用户在界面中看到的不是互不关联的零散记录，而是一条跨代理、跨任务、跨产物的连续历史线索。对于毕业论文这类强调过程的任务，这种设计比单独展示“最近聊天记录”更具有解释价值。",
                ],
                "figures": [
                    {
                        "caption": "图4-2 工作区活动时间线界面占位",
                        "hint": "建议插入 Workspace Activity Timeline 截图，突出 feature、chat、subagent 与 artifact 的统一时间线。",
                    }
                ],
            },
        ],
    },
    {
        "title": "第5章 核心功能模块与界面实现",
        "sections": [
            {
                "title": "5.1 工作区工作台与统一入口实现",
                "paragraphs": [
                    "在产品设计上，问津并没有把不同能力拆成彼此孤立的页面，而是围绕工作区建立统一工作台。用户先进入 workspace，再从其中触发 chat、feature 和 artifact follow-up。系统使用 `/workspaces/{workspace_id}/chat` 作为主对话入口，将 feature seed 参数附着到首条消息的 metadata 中，从而让编排意图可以在对话层自然延续。这种单工作区主线程模型减少了导航成本，也使 Harness 的“单入口控制面”能够在产品侧真实落地。",
                    "工作区主页中会根据 artifact、module status 和当前进度推导推荐的下一步 feature。对于已有深度调研结果但尚未写作的工作区，系统会优先推荐大纲或论文写作模块；对于已有章节草稿的工作区，则优先推荐评审、编译或导出能力。也就是说，界面层并不是静态功能展示，而是建立在 Summary 和 Artifact 状态之上的“可推进工作流”界面。",
                ],
                "figures": [
                    {
                        "caption": "图5-1 工作区主页/聊天入口截图占位",
                        "hint": "建议插入工作区主页或聊天页截图，展示单工作区主线程、模块卡片和右侧 Inspector。",
                    }
                ],
            },
            {
                "title": "5.2 深度调研与开题支撑模块实现",
                "paragraphs": [
                    "深度调研模块是毕业论文工作流中的前置能力。其 graph 会先读取 topic、discipline、focus_areas 和 memory_context，再组织 discovery、gap_mining、synthesis 和 cross_validation 四个阶段推进。第一阶段通过并行方式发现经典文献、近期工作和研究趋势；第二阶段识别研究空白；第三阶段综合为结构化结论；第四阶段做交叉校验。整个过程中，系统会持续输出 runtime block，例如发现摘要、研究空白列表和下一步建议动作。",
                    "该模块之所以能成为问津的核心，不在于它“能搜论文”，而在于它与后续写作模块天然联动。深度调研的输出可以直接转入文献管理、大纲生成和开题调研，形成前后呼应的工作链路。这种能力并不是某个 prompt 单独实现的，而是依赖于 Harness 对 feature graph、artifact 持久化和 summary 投影的统一支撑。",
                ],
            },
            {
                "title": "5.3 论文写作模块实现",
                "paragraphs": [
                    "论文写作模块通过 `thesis_writing_graph` 支持 `generate_outline`、`write_chapter`、`write_all`、`review_section` 和 `revise_section` 等多种动作。其核心思想不是一次生成整篇论文，而是将大纲、章节、评审和改写解耦为可组合能力。这样，用户可以围绕某一个阶段多次迭代，而不必在每次请求中重新生成整篇内容。",
                    "这一设计与 Memory System 和 Artifact System 形成了良好配合。一方面，长期记忆可以帮助系统保持用户偏好、写作风格和研究目标的一致性；另一方面，章节草稿和大纲作为 artifact 可被持续回收和再加工。因此，论文写作模块本身看似是“功能模块”，其背后实际运行的是完整的 Harness 协作链路。",
                ],
            },
            {
                "title": "5.4 Activity Timeline与Artifact Library实现",
                "paragraphs": [
                    "问津在前端实现中专门提供了 WorkspaceActivityTimeline 组件，用于展示 feature task、chat thread、subagent task 与 artifact 的统一时间线。用户可以通过筛选器按模块查看活动，并通过时间线项跳回对应的对话或结果详情。这一设计把复杂代理系统中的“后台过程”前置为用户可见对象，显著降低了系统黑箱感。",
                    "与此同时，Artifact Library 负责展示工作区中已经沉淀的结果，包括调研报告、章节草稿、图表与编译结果。由于 artifact 自身具有类型、版本与父子关系，系统可以在界面层进一步实现最近产物、follow-up 行为和 lineage 查看。对于毕业论文而言，这种“中间结果可见”能力极大改善了写作过程中的版本管理体验。",
                ],
                "figures": [
                    {
                        "caption": "图5-2 活动时间线与成果库截图占位",
                        "hint": "建议插入 Workspace Activity Timeline 或 Artifact Library 截图，突出统一时间线和产物沉淀。",
                    }
                ],
            },
            {
                "title": "5.5 LaTeX主稿台与编译导出实现",
                "paragraphs": [
                    "很多 AI 写作系统只能输出一段文本，剩余的排版和导出工作仍由用户手工完成。问津通过 LaTeXProjectService 与 CompileService 进一步向主稿交付推进。系统支持创建项目、复制模板、维护主文件、保存文件树排序、上传资源文件，并可在 Docker 化的 TeX Live 镜像中执行编译。编译结束后，系统会记录历史日志、PDF 路径、history_id 和错误信息，并提供 PDF 查看接口。",
                    "这一链路使问津不再是“聊天生成器”，而是“主稿生产工具”。对于工程开发型毕业论文而言，论文交付不仅包括正文文字，还包括图表、源文件、编译结果和后续修订能力。LaTeX 主稿台的存在，使平台在最终交付闭环上具备明显区别于普通对话系统的工程价值。",
                ],
                "figures": [
                    {
                        "caption": "图5-3 LaTeX主稿台截图占位",
                        "hint": "建议插入 LaTeX 文件树、编辑区和 PDF 预览同屏截图，突出主稿台闭环能力。",
                    }
                ],
            },
        ],
    },
    {
        "title": "第6章 系统测试与分析",
        "sections": [
            {
                "title": "6.1 测试目标与测试思路",
                "paragraphs": [
                    "由于本文重点在于平台架构与关键机制实现，因此测试目标主要围绕三类问题展开：第一，统一活动投影是否稳定，即 task、chat、artifact 和 subagent 能否被正确聚合；第二，论文写作主链路是否可用，即 thesis writing feature 能否完成端到端执行；第三，Super Agent Harness 的架构价值是否能够从测试和代码结构上得到印证。",
                    "本稿以自动化测试为主，重点验证核心后端链路。前端可视化截图和界面录屏则作为论文插图材料单独补充。这样的测试策略与毕业设计阶段目标一致，即优先证明系统架构是可运行、可验证、可继续扩展的。",
                ],
            },
            {
                "title": "6.2 自动化测试结果",
                "paragraphs": [
                    "本文实际运行了两组与平台主链路直接相关的测试。第一组为 `workspace_activity_service` 测试，执行命令为 `uv run pytest tests/services/test_workspace_activity_service.py -q`，测试结果为 7 passed in 0.73s。第二组为 `thesis_writing_end_to_end` 测试，执行命令为 `uv run pytest tests/task/test_thesis_writing_end_to_end.py -q`，测试结果为 3 passed in 0.62s。",
                    "这两组测试虽然数量有限，但覆盖面具有代表性。前者验证了统一活动投影机制，后者验证了论文写作 feature 从输入 payload 到结构化结果的主链路。它们共同说明，问津并不是只在概念层面提出了 Super Agent Harness，而是已经在关键路径上具备实际运行能力。",
                ],
                "table": {
                    "title": "表6-1 已执行自动化测试结果",
                    "headers": ["测试项", "执行命令", "结果", "说明"],
                    "rows": [
                        ["工作区活动聚合测试", "uv run pytest tests/services/test_workspace_activity_service.py -q", "7 passed in 0.73s", "验证 task/chat/subagent/artifact 聚合逻辑"],
                        ["论文写作端到端测试", "uv run pytest tests/task/test_thesis_writing_end_to_end.py -q", "3 passed in 0.62s", "验证 thesis writing 主链路可用性"],
                    ],
                },
            },
            {
                "title": "6.3 架构设计效果分析",
                "paragraphs": [
                    "结合系统实现与测试结果，可以从三个角度分析架构效果。首先是统一性。无论是深度调研、论文写作还是图表生成，系统都围绕同一 feature execute 入口、同一 task runtime 和同一 result contract 运转，这降低了功能扩张时的维护成本。其次是连续性。Memory System 与 ThreadState 使用户偏好、研究阶段和长期目标能够在多个回合之间被保留和重用，解决了传统 AI 对话易失上下文的问题。最后是可观测性。Activity Timeline 与 Artifact Library 让系统执行过程可以被用户直接看到，从而降低复杂代理系统的黑箱感。",
                    "从毕业论文场景看，以上三点非常关键。因为毕业论文不是单次对话问题，而是需要经历大量小步骤、长期反复修改和多轮导师反馈的过程。问津通过 Harness 把这些过程纳入统一控制面，使平台更像一个真正的写作工作台，而不是若干 AI 能力的松散集合。",
                ],
            },
            {
                "title": "6.4 当前不足与改进方向",
                "paragraphs": [
                    "尽管当前系统已经具备较完整的架构骨架，但仍存在进一步优化空间。首先，测试覆盖范围仍需扩展，尤其是前端交互、LaTeX 编译联调、Docker 异常恢复和更复杂的 subagent 并发场景。其次，长期记忆虽然已经具备注入、抽取和压缩机制，但未来仍可进一步增强冲突检测、时效性管理和更细粒度的用户控制。再次，当前对引用治理、事实校验和学术规范检查的支撑还偏轻量，后续可以通过更强的 reviewer/validator 机制进行补足。",
                    "这些问题并不否定本文方案的有效性，反而说明问津已经从“是否能够搭建”进入“如何持续完善”的阶段。对于毕业设计项目而言，完成统一架构、关键机制和代表性测试，已经证明本系统具备继续演化的工程基础。",
                ],
            },
        ],
    },
    {
        "title": "第7章 总结与展望",
        "sections": [
            {
                "title": "7.1 研究总结",
                "paragraphs": [
                    "本文围绕问津学术写作工作台，提出并实现了以 Super Agent Harness 为核心的平台架构，并将长期记忆系统作为其中的重要基础能力。不同于只强调文本生成功能的传统系统，本文更关注多智能体学术写作平台如何建立清晰的架构边界、统一的执行主链路和稳定的长期上下文能力。",
                    "在实现层面，本文完成了平台分层设计、canonical route 与 result contract 设计、Lead Agent 与中间件链设计、Memory System 设计、Subagent 调度设计、Task Runtime 设计，以及基于 Artifact 与 Activity 的过程投影设计。系统还围绕毕业论文工作区提供了深度调研、论文写作、工作区时间线与 LaTeX 主稿台等关键模块，并通过自动化测试验证了代表性功能链路。",
                    "因此，本文认为，学术写作平台的核心竞争力不仅在于模型是否足够强，更在于是否拥有能够稳定组织模型能力的软件底座。Super Agent Harness 正是问津平台的这一底座，而 Memory System 则是其支撑长期写作连续性的关键能力。",
                ],
            },
            {
                "title": "7.2 后续展望",
                "paragraphs": [
                    "未来工作可以从四个方向继续推进。第一，进一步增强 Memory System 的冲突解决、时间衰减和用户可控性，使长期记忆更适合更长周期的研究协作。第二，继续扩展 Harness 的可观测性，例如引入 phase 级耗时统计、模型成本统计和异常类型分析。第三，强化 reviewer、citation validator 和 fact checker 等专业子代理，使平台在学术规范方面具备更强支撑。第四，增强 LaTeX 主稿与章节 artifact 之间的双向映射能力，让论文写作与最终版式编辑更加紧密地协同。",
                    "总体而言，问津的实践说明，下一代学术写作平台的关键不在于把多少模型接入系统，而在于能否用统一架构把这些能力组织成一个可持续演进的软件系统。围绕 Super Agent Harness 和 Memory System 的持续深化，将是该方向最值得投入的工程路径之一。",
                ],
            },
        ],
    },
]


REFERENCES = [
    "[1] Yao S, Zhao J, Yu D, et al. ReAct: Synergizing reasoning and acting in language models[EB/OL]. arXiv:2210.03629, 2022.",
    "[2] Lewis P, Perez E, Piktus A, et al. Retrieval-augmented generation for knowledge-intensive NLP tasks[C]//Advances in Neural Information Processing Systems. 2020.",
    "[3] Wei J, Wang X, Schuurmans D, et al. Chain-of-thought prompting elicits reasoning in large language models[C]//Advances in Neural Information Processing Systems. 2022.",
    "[4] Ramírez S. FastAPI documentation[EB/OL]. https://fastapi.tiangolo.com/, 2026-04-08.",
    "[5] Vercel. Next.js documentation[EB/OL]. https://nextjs.org/docs, 2026-04-08.",
    "[6] Bayer M. SQLAlchemy 2.0 documentation[EB/OL]. https://docs.sqlalchemy.org/, 2026-04-08.",
    "[7] The PostgreSQL Global Development Group. PostgreSQL documentation[EB/OL]. https://www.postgresql.org/docs/, 2026-04-08.",
    "[8] Redis Ltd. Redis documentation[EB/OL]. https://redis.io/docs/, 2026-04-08.",
    "[9] LangChain, Inc. LangGraph documentation[EB/OL]. https://docs.langchain.com/oss/python/langgraph/overview, 2026-04-08.",
    "[10] Docker, Inc. Docker documentation[EB/OL]. https://docs.docker.com/, 2026-04-08.",
    "[11] Vaswani A, Shazeer N, Parmar N, et al. Attention Is All You Need[EB/OL]. arXiv:1706.03762, 2017.",
    "[12] Brown T B, Mann B, Ryder N, et al. Language Models are Few-Shot Learners[EB/OL]. arXiv:2005.14165, 2020.",
    "[13] OpenAI, Achiam J, Adler S, et al. GPT-4 Technical Report[EB/OL]. arXiv:2303.08774, 2023.",
    "[14] Touvron H, Martin L, Stone K, et al. Llama 2: Open Foundation and Fine-Tuned Chat Models[EB/OL]. arXiv:2307.09288, 2023.",
    "[15] Schick T, Dwivedi-Yu J, Dessì R, et al. Toolformer: Language Models Can Teach Themselves to Use Tools[EB/OL]. arXiv:2302.04761, 2023.",
    "[16] Madaan A, Tandon N, Gupta P, et al. Self-Refine: Iterative Refinement with Self-Feedback[EB/OL]. arXiv:2303.17651, 2023.",
    "[17] Shinn N, Cassano F, Berman E, et al. Reflexion: Language Agents with Verbal Reinforcement Learning[EB/OL]. arXiv:2303.11366, 2023.",
    "[18] Yao S, Yu D, Zhao J, et al. Tree of Thoughts: Deliberate Problem Solving with Large Language Models[EB/OL]. arXiv:2305.10601, 2023.",
    "[19] Wu Q, Bansal G, Zhang J, et al. AutoGen: Enabling Next-Gen LLM Applications via Multi-Agent Conversation[EB/OL]. arXiv:2308.08155, 2023.",
    "[20] Hong S, Zhuge M, Chen J, et al. MetaGPT: Meta Programming for A Multi-Agent Collaborative Framework[EB/OL]. arXiv:2308.00352, 2023.",
    "[21] Li G, Hammoud H A A K, Itani H, et al. CAMEL: Communicative Agents for \"Mind\" Exploration of Large Language Model Society[EB/OL]. arXiv:2303.17760, 2023.",
    "[22] Wang L, Ma C, Feng X, et al. A Survey on Large Language Model based Autonomous Agents[EB/OL]. arXiv:2308.11432, 2023.",
    "[23] Park J S, O'Brien J C, Cai C J, et al. Generative Agents: Interactive Simulacra of Human Behavior[EB/OL]. arXiv:2304.03442, 2023.",
    "[24] Packer C, Wooders S, Lin K, et al. MemGPT: Towards LLMs as Operating Systems[EB/OL]. arXiv:2310.08560, 2023.",
    "[25] Zhong W, Guo L, Gao Q, et al. MemoryBank: Enhancing Large Language Models with Long-Term Memory[EB/OL]. arXiv:2305.10250, 2023.",
    "[26] Wang W, Dong L, Cheng H, et al. Augmenting Language Models with Long-Term Memory[EB/OL]. arXiv:2306.07174, 2023.",
    "[27] Wang G, Xie Y, Jiang Y, et al. Voyager: An Open-Ended Embodied Agent with Large Language Models[EB/OL]. arXiv:2305.16291, 2023.",
    "[28] Celery Project. User Guide[EB/OL]. https://docs.celeryq.dev/en/stable/userguide/index.html, 2026-04-08.",
    "[29] Meta Platforms, Inc. React documentation[EB/OL]. https://react.dev/, 2026-04-08.",
    "[30] Microsoft. TypeScript documentation[EB/OL]. https://www.typescriptlang.org/docs/, 2026-04-08.",
]


ACKNOWLEDGEMENT = [
    "在本论文架构强化版初稿完成过程中，首先要感谢指导教师和学院老师在选题方向、系统设计思路和论文撰写规范方面给予的指导。本稿仍有封面字段、正式截图与个别图表待进一步完善，但整体结构与主要论证框架已经基本形成。",
    "同时，感谢问津项目现有代码、文档和测试基础，为本文提供了系统化的工程材料。也感谢在毕业设计阶段给予建议和反馈的同学与朋友，使本文能够围绕 Super Agent Harness 与 Memory System 形成较为完整的总结。",
]


def set_doc_defaults(doc: Document) -> None:
    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(3.17)
    section.right_margin = Cm(3.17)
    section.header_distance = Cm(1.5)
    section.footer_distance = Cm(1.75)

    normal = doc.styles["Normal"]
    normal.font.name = "Times New Roman"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    normal.font.size = Pt(12)


def set_run_font(
    run,
    *,
    east_asia: str = "宋体",
    latin: str = "Times New Roman",
    size: int = 12,
    bold: bool = False,
) -> None:
    run.font.name = latin
    run._element.rPr.rFonts.set(qn("w:eastAsia"), east_asia)
    run.font.size = Pt(size)
    run.bold = bold


def format_paragraph(
    paragraph,
    *,
    align=WD_ALIGN_PARAGRAPH.JUSTIFY,
    first_line: float = 0.74,
    line_spacing: float = 1.5,
    space_before: int = 0,
    space_after: int = 0,
) -> None:
    paragraph.alignment = align
    paragraph.paragraph_format.first_line_indent = Cm(first_line) if first_line else Cm(0)
    paragraph.paragraph_format.line_spacing = line_spacing
    paragraph.paragraph_format.space_before = Pt(space_before)
    paragraph.paragraph_format.space_after = Pt(space_after)


def add_body_paragraph(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    format_paragraph(p)
    run = p.add_run(text)
    set_run_font(run)


def add_center_paragraph(
    doc: Document,
    text: str,
    *,
    size: int = 16,
    bold: bool = True,
    east_asia: str = "黑体",
) -> None:
    p = doc.add_paragraph()
    format_paragraph(
        p,
        align=WD_ALIGN_PARAGRAPH.CENTER,
        first_line=0,
        line_spacing=1.5,
        space_before=6,
        space_after=6,
    )
    run = p.add_run(text)
    set_run_font(run, east_asia=east_asia, size=size, bold=bold)


def add_subheading(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    format_paragraph(
        p,
        align=WD_ALIGN_PARAGRAPH.LEFT,
        first_line=0,
        line_spacing=1.5,
        space_before=6,
        space_after=6,
    )
    run = p.add_run(text)
    set_run_font(run, east_asia="黑体", size=14, bold=True)


def add_keyword_line(doc: Document, label: str, content: str) -> None:
    p = doc.add_paragraph()
    format_paragraph(p, align=WD_ALIGN_PARAGRAPH.LEFT, first_line=0, line_spacing=1.5)
    run1 = p.add_run(label)
    set_run_font(run1, east_asia="黑体", size=12, bold=True)
    run2 = p.add_run(content)
    set_run_font(run2, size=12)


def add_table(doc: Document, title: str, headers: list[str], rows: list[list[str]]) -> None:
    add_center_paragraph(doc, title, size=12, bold=False, east_asia="宋体")
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    try:
        table.style = "Table Grid"
    except KeyError:
        pass

    hdr_cells = table.rows[0].cells
    for idx, header in enumerate(headers):
        hdr_cells[idx].text = header
        hdr_cells[idx].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        p = hdr_cells[idx].paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in p.runs:
            set_run_font(run, east_asia="黑体", size=10, bold=True)

    for row in rows:
        cells = table.add_row().cells
        for idx, value in enumerate(row):
            cells[idx].text = value
            cells[idx].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            for p in cells[idx].paragraphs:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER if idx == 0 else WD_ALIGN_PARAGRAPH.LEFT
                p.paragraph_format.line_spacing = 1.2
                for run in p.runs:
                    set_run_font(run, size=10)
    doc.add_paragraph()


def add_figure_placeholder(doc: Document, caption: str, hint: str) -> None:
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    try:
        table.style = "Table Grid"
    except KeyError:
        pass
    cell = table.cell(0, 0)
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    p1 = cell.paragraphs[0]
    p1.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r1 = p1.add_run("【截图/示意图占位】")
    set_run_font(r1, east_asia="黑体", size=12, bold=True)

    p2 = cell.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r2 = p2.add_run(hint)
    set_run_font(r2, size=11)

    for _ in range(3):
        p = cell.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(" ")
        set_run_font(r, size=11)

    cap = doc.add_paragraph()
    format_paragraph(
        cap,
        align=WD_ALIGN_PARAGRAPH.CENTER,
        first_line=0,
        line_spacing=1.2,
        space_before=3,
        space_after=6,
    )
    run = cap.add_run(caption)
    set_run_font(run, size=11)


def add_page_break(doc: Document) -> None:
    p = doc.add_paragraph()
    run = p.add_run()
    run.add_break(WD_BREAK.PAGE)


def fill_cover(doc: Document, meta: CoverMeta) -> None:
    table = doc.tables[0]
    table.cell(0, 1).text = meta.title
    table.cell(1, 1).text = meta.college
    table.cell(2, 1).text = meta.major
    table.cell(3, 1).text = meta.class_name
    table.cell(4, 1).text = meta.student_id
    table.cell(5, 1).text = meta.student_name
    table.cell(6, 1).text = meta.advisor
    table.cell(7, 1).text = meta.advisor_title
    table.cell(8, 1).text = meta.completion_date

    for row in table.rows:
        for cell in row.cells:
            for p in cell.paragraphs:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER if p is cell.paragraphs[0] else WD_ALIGN_PARAGRAPH.LEFT
                for run in p.runs:
                    set_run_font(run, east_asia="宋体", size=12)


def add_toc(doc: Document) -> None:
    add_page_break(doc)
    add_center_paragraph(doc, "目    录", size=16, bold=True, east_asia="黑体")
    for item in TOC_ITEMS:
        p = doc.add_paragraph()
        format_paragraph(p, align=WD_ALIGN_PARAGRAPH.LEFT, first_line=0, line_spacing=1.5)
        run = p.add_run(item)
        set_run_font(run, size=12)


def add_abstract(doc: Document) -> None:
    add_page_break(doc)
    add_center_paragraph(doc, "摘    要", size=16, bold=True, east_asia="黑体")
    for paragraph in ABSTRACT_ZH:
        add_body_paragraph(doc, paragraph)
    add_keyword_line(doc, "关键词：", KEYWORDS_ZH)

    add_page_break(doc)
    add_center_paragraph(doc, "ABSTRACT", size=16, bold=True, east_asia="黑体")
    for paragraph in ABSTRACT_EN:
        add_body_paragraph(doc, paragraph)
    add_keyword_line(doc, "Key words: ", KEYWORDS_EN)


def add_body(doc: Document) -> None:
    add_page_break(doc)
    for chapter in BODY:
        add_center_paragraph(doc, chapter["title"], size=16, bold=True, east_asia="黑体")
        for section in chapter["sections"]:
            add_subheading(doc, section["title"])
            for paragraph in section["paragraphs"]:
                add_body_paragraph(doc, paragraph)
            table = section.get("table")
            if table:
                add_table(doc, table["title"], table["headers"], table["rows"])
            for figure in section.get("figures", []):
                add_figure_placeholder(doc, figure["caption"], figure["hint"])


def add_references(doc: Document) -> None:
    add_center_paragraph(doc, "参考文献", size=16, bold=True, east_asia="黑体")
    for item in REFERENCES:
        p = doc.add_paragraph()
        format_paragraph(p, align=WD_ALIGN_PARAGRAPH.LEFT, first_line=0, line_spacing=1.5)
        run = p.add_run(item)
        set_run_font(run, size=12)


def add_acknowledgement(doc: Document) -> None:
    add_center_paragraph(doc, "致谢", size=16, bold=True, east_asia="黑体")
    for paragraph in ACKNOWLEDGEMENT:
        add_body_paragraph(doc, paragraph)


def ensure_template_exists() -> None:
    if not TEMPLATE.exists():
        raise FileNotFoundError(f"Template not found: {TEMPLATE}")


def build_document() -> Path:
    ensure_template_exists()
    doc = Document(str(TEMPLATE))
    set_doc_defaults(doc)
    fill_cover(doc, COVER)
    add_toc(doc)
    add_abstract(doc)
    add_body(doc)
    add_references(doc)
    add_acknowledgement(doc)
    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(OUT))
    return OUT


def main() -> None:
    out = build_document()
    print(out)


if __name__ == "__main__":
    main()
